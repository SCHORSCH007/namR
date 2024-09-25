import docx

def docx_strategy(file_path, wrapper: str="only return an adequate filename to the following docx-file", description:str=None):
    text = read_docx(file_path=file_path)
    metadata = read_docx_metadata(file_path=file_path)
    
    first_words = ""
    metadata_str = ""
    description_str = ""
    
    if description:
        description_str = f"\n description what the file is: ```{description}```"
    if metadata:
        metadata_str = ' '.join([f"{k}: {v}" for k, v in metadata.items()])
    if text:
        first_words = ' '.join(text)
    
    string_request = f"{wrapper} {description_str} \n metadata: ```{metadata_str}``` \n first words: ```{first_words}``` \n only return an adequate filename"
    
    return string_request



def read_docx_metadata(file_path):
    try:
        doc = docx.Document(file_path)
        metadata = {}

        # Access core properties
        core_properties = doc.core_properties
        metadata['title'] = core_properties.title
        metadata['subject'] = core_properties.subject
        metadata['keywords'] = core_properties.keywords
        metadata['creation_date'] = core_properties.created
        metadata['modified_date'] = core_properties.modified

        return {k: v for k,v in metadata.items() if v }

    except Exception as e:
        print(f"Error reading metadata: {e} in file {file_path}")
        return None
    
def read_docx (file_path: str, count: int=100):
    try:
        doc = docx.Document(file_path)
        
        text = []
        for para in doc.paragraphs:

            words = para.text.split()

            count = count - len(words)
            if count <= 0:
                text.extend(words[:(len(words)+count)])
                return text
            else:
                text.extend(words)
            

    except Exception as e:
        print(f"failed reading docx {file_path}")
        return []